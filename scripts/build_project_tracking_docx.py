from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent.parent
OUT_DIR = ROOT / "docs" / "01_final_report"
OUT_PATH = OUT_DIR / "WasteWise_Project_Tracking_Report.docx"


BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "0B2545"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
CALLOUT = "F4F6F9"
GREEN = "1F7A3A"
GOLD = "7A5A00"
RED = "9B1C1C"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths_dxa: list[int]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[min(idx, len(widths_dxa) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_cell_text(cell, text: str, bold=False, color=None, size=9.5) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def style_table(table, widths_dxa: list[int], header_fill=LIGHT_BLUE) -> None:
    table.style = "Table Grid"
    set_table_width(table, widths_dxa)
    set_repeat_table_header(table.rows[0])
    for row_i, row in enumerate(table.rows):
        for cell in row.cells:
            if row_i == 0:
                set_cell_shading(cell, header_fill)
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.bold = True
                        r.font.color.rgb = RGBColor.from_string(INK)
            elif row_i % 2 == 0:
                set_cell_shading(cell, "FBFCFD")


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[int]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    for i, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], header, bold=True, color=INK, size=9)
    for row_values in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row_values):
            set_cell_text(cells[i], str(value), size=9)
    style_table(table, widths)
    doc.add_paragraph()


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.add_run(item)


def add_callout(doc: Document, title: str, body: str, fill=CALLOUT) -> None:
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, top=130, bottom=130, start=160, end=160)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    r.font.size = Pt(11)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.15
    p2.add_run(body).font.size = Pt(10)
    style_table(table, [9360], header_fill=fill)
    doc.add_paragraph()


def add_figure(doc: Document, path: Path, caption: str, width_in=5.9) -> None:
    if not path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width_in))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(8)
    r = cap.add_run(caption)
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor.from_string("555555")


def configure_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    title = styles["Title"]
    title.font.name = "Calibri"
    title.font.size = Pt(24)
    title.font.color.rgb = RGBColor.from_string(INK)
    title.paragraph_format.space_after = Pt(6)

    subtitle = styles["Subtitle"]
    subtitle.font.name = "Calibri"
    subtitle.font.size = Pt(12)
    subtitle.font.color.rgb = RGBColor.from_string("555555")
    subtitle.paragraph_format.space_after = Pt(12)

    for style_name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        st = styles[style_name]
        st.font.name = "Calibri"
        st.font.size = Pt(size)
        st.font.color.rgb = RGBColor.from_string(color)
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True

    for style_name in ["List Bullet", "List Number"]:
        st = styles[style_name]
        st.font.name = "Calibri"
        st.font.size = Pt(11)
        st.paragraph_format.left_indent = Inches(0.375)
        st.paragraph_format.first_line_indent = Inches(-0.188)
        st.paragraph_format.space_after = Pt(4)
        st.paragraph_format.line_spacing = 1.25


def add_header_footer(doc: Document) -> None:
    section = doc.sections[0]
    header_p = section.header.paragraphs[0]
    header_p.text = "WasteWise FYP Project Tracking Report"
    header_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for r in header_p.runs:
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor.from_string("666666")

    footer_p = section.footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_p.add_run("Generated: 2026-05-30")
    for r in footer_p.runs:
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor.from_string("666666")


def build_doc() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_styles(doc)
    add_header_footer(doc)

    doc.add_paragraph("WasteWise: Project Status and Experiment Tracking Report", style="Title")
    doc.add_paragraph("Final Year Project: automated waste classification, feature-based ML, and localization-first classification", style="Subtitle")
    add_callout(
        doc,
        "Current project position",
        "Current datasets are `external_datasets/super_yolo_dataset` for YOLO localization and `data/merged_dataset_v5` for classification. A newest-dataset 637-feature ML rerun now exists, while older `merged_dataset_v3` results remain useful historical lecturer evidence.",
        fill="EAF4EE",
    )

    add_table(
        doc,
        ["Area", "Current status", "Main evidence"],
        [
            ["Dataset", "Use newest datasets for current tracking.", "super_yolo_dataset: 23,929 images / 102,777 boxes; merged_dataset_v5: 29,639 classification images."],
            ["ML", "Keep as main explainable pipeline; report legacy and newest-dataset evidence separately.", "Newest super_yolo_dataset rerun: XGBoost accuracy 0.5408, F1-macro 0.3691. Legacy lecturer artifact: 0.6742 / 0.6506."],
            ["PCA", "Controlled model sweep completed.", "637 -> 128 keeps 99.90% variance; Linear SVM drops 62.43% -> 59.90%, a 2.52 pp trade-off."],
            ["DL classification", "Use as comparison / gate, not final localization metric.", "EfficientNetB0 crop classifier 94.29% in architecture comparison."],
            ["DL localization", "Improved after Stage 2 rework.", "YOLO localization-only conf=0.30 on 300 images: precision 0.6999, recall 0.5729, F1 0.6301, mean IoU 0.9057."],
        ],
        [1500, 4200, 3660],
    )

    doc.add_paragraph("1. Project Goal And Scope", style="Heading 1")
    doc.add_paragraph(
        "WasteWise is an FYP system for automated waste recognition and sorting. The work now separates explainable Machine Learning evidence from Deep Learning localization evidence, so each branch can be reported honestly and evaluated with the right metric."
    )
    add_bullets(
        doc,
        [
            "Primary ML goal: classify waste object crops using lecturer-explainable handcrafted features.",
            "Primary DL goal: perform YOLO object localization first, then verify/classify crops.",
            "Deployment direction: web-based inference with sorting guidance, history, and settings.",
        ],
    )

    doc.add_paragraph("2. Dataset Tracking", style="Heading 1")
    doc.add_paragraph(
        "Newest datasets detected in the workspace are listed below. These should be treated as the current dataset sources for the report. Older dataset names are retained only when a specific saved experiment was produced from them."
    )
    add_table(
        doc,
        ["YOLO class", "Train boxes", "Val boxes", "Test boxes", "Total boxes"],
        [
            ["plastic", "18,418", "1,568", "1,668", "21,654"],
            ["glass", "7,186", "2,488", "9", "9,683"],
            ["metal", "8,785", "1,800", "542", "11,127"],
            ["paper", "4,731", "188", "1,347", "6,266"],
            ["cardboard", "7,433", "1,614", "35", "9,082"],
            ["organic", "32,171", "12,748", "46", "44,965"],
            ["Total boxes", "78,724", "20,406", "3,647", "102,777"],
        ],
        [1700, 1900, 1900, 1900, 1960],
    )
    add_table(
        doc,
        ["Dataset", "Split", "Images", "Notes"],
        [
            ["external_datasets/super_yolo_dataset", "train", "19,559", "YOLO-format images with label files."],
            ["external_datasets/super_yolo_dataset", "val", "3,266", "YOLO-format images with label files."],
            ["external_datasets/super_yolo_dataset", "test", "1,104", "YOLO-format images with label files."],
            ["data/merged_dataset_v5", "train", "24,039", "Classification-folder dataset; includes Background."],
            ["data/merged_dataset_v5", "test", "5,600", "Balanced 800 images per class."],
        ],
        [3100, 1100, 1200, 3960],
    )
    add_table(
        doc,
        ["merged_dataset_v5 class", "Train images", "Test images", "Total"],
        [
            ["Background", "3,500", "800", "4,300"],
            ["cardboard", "3,425", "800", "4,225"],
            ["glass", "3,500", "800", "4,300"],
            ["metal", "3,230", "800", "4,030"],
            ["organic", "3,500", "800", "4,300"],
            ["paper", "3,500", "800", "4,300"],
            ["plastic", "3,384", "800", "4,184"],
            ["Total", "24,039", "5,600", "29,639"],
        ],
        [3100, 2000, 2000, 2260],
    )
    add_callout(
        doc,
        "Reporting note",
        "The newest YOLO dataset is imbalanced, especially in validation/test class distribution. The newest classification dataset is much more balanced and includes a Background class. Legacy ML metrics from `merged_dataset_v3` are now labelled separately from the newest `super_yolo_dataset` ML rerun.",
        fill="FFF7DF",
    )

    doc.add_paragraph("3. Machine Learning Pipeline", style="Heading 1")
    doc.add_paragraph(
        "The ML branch is the most explainable and stable part of the project. It uses object crops from YOLO labels, resizes crops internally, extracts a fixed 637-D vector, then compares classical ML models."
    )
    add_callout(
        doc,
        "Dataset alignment note",
        "There are now two ML evidence tracks: the stronger lecturer-facing legacy run on `merged_dataset_v3`, and a newest-dataset rerun on `super_yolo_dataset`. Use the newest run when answering dataset-version questions; use the legacy run as historical evidence only.",
        fill="FFF7DF",
    )
    add_table(
        doc,
        ["Feature group", "Count", "Purpose"],
        [
            ["Spatial", "8", "Intensity, gradients, edge density."],
            ["Frequency / FFT", "9", "Radial frequency energy and high-frequency texture."],
            ["Color", "44", "HSV histograms plus BGR/HSV mean and standard deviation."],
            ["HOG", "576", "Local shape and gradient-orientation texture."],
            ["Total", "637", "Fixed handcrafted representation for explainable ML."],
        ],
        [2200, 1200, 5960],
    )
    add_table(
        doc,
        ["Current-dataset model", "Accuracy", "F1-macro", "Decision"],
        [
            ["XGBoost", "0.5408", "0.3691", "Best on newest `super_yolo_dataset` rerun."],
            ["Random Forest", "0.5063", "0.3456", "Tree baseline with feature importance chart."],
            ["ExtraTrees", "0.5045", "0.3414", "Close to RF on newest data."],
            ["Linear SVM", "0.4628", "0.3159", "Margin-based baseline."],
            ["Logistic Regression", "0.4494", "0.3054", "Linear standardized baseline."],
            ["Decision Tree", "0.3750", "0.2631", "Simple interpretable baseline."],
        ],
        [2500, 1400, 1400, 4060],
    )
    add_callout(
        doc,
        "Newest ML support warning",
        "The newest ML rerun used 24,000 train crops, but only 2,232 test crops because the test split has very low support for glass (9), cardboard (35), and organic (46). Report the lower F1 as current-dataset evidence under test imbalance, not as a balanced benchmark.",
        fill="FFF7DF",
    )
    add_table(
        doc,
        ["Legacy lecturer model", "Accuracy", "F1-macro", "Decision"],
        [
            ["XGBoost", "0.6742", "0.6506", "Best lecturer-facing ML result."],
            ["Random Forest", "0.6317", "0.6111", "Useful for feature importance."],
            ["ExtraTrees", "0.6312", "0.6113", "Strong high-dimensional tree baseline."],
            ["Linear SVM", "0.5960", "0.5642", "Margin-based baseline."],
            ["Logistic Regression", "0.5864", "0.5558", "Linear standardized baseline."],
            ["Decision Tree", "0.5115", "0.4883", "Simple interpretable baseline."],
        ],
        [2100, 1500, 1500, 4260],
    )
    add_table(
        doc,
        ["Feature importance group", "Importance"],
        [
            ["HOG", "59.5808%"],
            ["Color", "29.2090%"],
            ["Frequency", "5.6582%"],
            ["Spatial", "5.5520%"],
        ],
        [5200, 4160],
    )
    add_figure(doc, ROOT / "runs" / "ml" / "feature_ml_super_yolo_6class_4k" / "chart_model_comparison.png", "Figure 1. Current-dataset classical ML model comparison.", 5.8)
    add_figure(doc, ROOT / "runs" / "ml" / "feature_ml_super_yolo_6class_4k" / "chart_domain_importance.png", "Figure 2. Current-dataset feature group / domain importance.", 5.8)

    doc.add_paragraph("4. PCA Dimensionality Reduction", style="Heading 1")
    add_table(
        doc,
        ["Evidence", "Model", "Components", "Accuracy", "F1", "Drop"],
        [
            ["Controlled ML sweep", "Linear SVM", "637", "62.43%", "0.6235", "0.00 pp"],
            ["Controlled ML sweep", "Linear SVM", "128", "59.90%", "0.5947", "2.52 pp"],
            ["Controlled ML sweep", "Logistic Regression", "637", "60.24%", "0.6019", "0.00 pp"],
            ["Controlled ML sweep", "Logistic Regression", "128", "59.71%", "0.5954", "0.52 pp"],
            ["ANN-only artifact", "MLP", "637", "73.24%", "0.7319", "0.00 pp"],
            ["ANN-only artifact", "MLP", "128", "68.71%", "0.6863", "4.53 pp"],
        ],
        [2300, 2100, 1300, 1300, 1200, 1260],
    )
    add_callout(
        doc,
        "PCA claim control",
        "Use the controlled ML sweep for the 'about 2%' statement: Linear SVM drops 2.52 percentage points when reducing 637 features to 128 PCA components. The older ANN-only PCA artifact drops 4.53 points, so it must be labelled separately.",
        fill="EAF4EE",
    )
    add_figure(doc, ROOT / "runs" / "ml" / "pca_feature_model_sweep" / "pca_model_sweep_accuracy.png", "Figure 3. PCA components vs accuracy by classical model.", 5.8)

    doc.add_paragraph("5. Deep Learning Experiments", style="Heading 1")
    add_table(
        doc,
        ["Experiment", "Result", "Use in final report"],
        [
            ["ANN/CNN crop baselines", "Tuned ANN accuracy 0.4057; tuned CNN accuracy 0.4413.", "Baseline evidence only."],
            ["CNN + ANN soft voting", "50/50 ensemble accuracy 78.69%, macro F1 78.12%.", "Shows feature complementarity."],
            ["Architecture comparison", "EfficientNetB0 94.29%; ResNet50 89.76%; MobileNetV2 85.43%.", "Supports EfficientNetB0 as classifier/gate."],
            ["Old 2-stage pipeline", "YOLO localization -> EfficientNetB0 verification; 100-image sweep had 295 accepted from 348 proposals.", "Experimental evidence; not final direction."],
        ],
        [2500, 3700, 3160],
    )
    add_table(
        doc,
        ["Deep model", "Accuracy", "Size", "Latency"],
        [
            ["MobileNetV2", "85.43%", "20.07 MB", "253.5 ms"],
            ["ResNet50", "89.76%", "161.52 MB", "163.1 ms"],
            ["EfficientNetB0", "94.29%", "29.21 MB", "288.6 ms"],
        ],
        [2500, 2100, 2300, 2460],
    )
    add_figure(doc, ROOT / "runs" / "dl" / "comparison_models" / "confusion_matrix_grid.png", "Figure 4. Deep model confusion matrix grid.", 5.8)

    doc.add_paragraph("6. Final DL Rework: Classification First, Localization Second", style="Heading 1")
    doc.add_paragraph(
        "The new DL requirement is reversed from the old YOLO-first pipeline. Stage 1 classifies or gates the image first. Stage 2 performs localization only. The class decision and box localization are tracked separately."
    )
    add_table(
        doc,
        ["Stage", "Role", "Current implementation"],
        [
            ["Stage 1", "Classification / image-level gate", "EfficientNetB0 classifier predicts image/crop class and diagnostic confidence."],
            ["Stage 2", "Localization only", "YOLO model is used only to output bounding boxes; YOLO class is not the final classifier."],
            ["Evaluation", "Localization metrics", "IoU@0.5 matching, precision, recall, mean matched IoU."],
        ],
        [1500, 2900, 4960],
    )
    add_table(
        doc,
        ["Stage 2 localizer", "Images", "Precision", "Recall", "F1", "Mean IoU", "TP/FP/FN"],
        [
            ["Grad-CAM baseline", "60", "0.2568", "0.0728", "0.1134", "0.7127", "19 / 55 / 242"],
            ["YOLO conf=0.25 ablation", "60", "0.6352", "0.5670", "0.5991", "0.9012", "148 / 85 / 113"],
            ["YOLO conf=0.30 promoted final", "300", "0.6999", "0.5729", "0.6301", "0.9057", "660 / 283 / 492"],
            ["YOLO conf=0.35 balanced sweep", "300", "0.7571", "0.5356", "0.6274", "0.9043", "617 / 198 / 535"],
            ["YOLO conf=0.40 precision sweep", "300", "0.8035", "0.5148", "0.6275", "0.9050", "593 / 145 / 559"],
        ],
        [3000, 850, 1000, 1000, 850, 1200, 1460],
    )
    add_callout(
        doc,
        "Recommended DL localization setting",
        "`--localizer yolo --yolo-conf 0.30` is the promoted balanced setting from the 300-image sweep: best F1/recall among tested thresholds and mean matched IoU above 0.90. Use `conf=0.40` only when the report wants the highest precision trade-off.",
        fill="EAF4EE",
    )
    add_figure(
        doc,
        ROOT / "runs" / "dl" / "localization_rework" / "yolo_conf030_stratified300_final" / "visuals" / "rf_garbage_metal391_jpg.rf.d2d79150c42df8cd64bea8d65acc58ab_yolo.jpg",
        "Figure 5. Example localization-first crop-verification output.",
        5.8,
    )

    doc.add_paragraph("7. Reproducible Commands", style="Heading 1")
    add_table(
        doc,
        ["Task", "Command / path"],
        [
            ["Current feature ML run", r".\.venv311\Scripts\python.exe scripts\feature_ml_analysis.py --data external_datasets\super_yolo_dataset\data.yaml --out runs\ml\feature_ml_super_yolo_6class_4k --exclude-classes= --max-per-class-train 4000 --max-per-class-test 800 --domain-out runs\ml\feature_ml_super_yolo_6class_4k\frequency_analysis"],
            ["Legacy feature ML run", r".\.venv311\Scripts\python.exe scripts\feature_ml_analysis.py --data merged_dataset_v3\data.yaml --out runs\ml\feature_ml_lecturer_6class_4k --exclude-classes other --max-per-class-train 4000 --max-per-class-test 800"],
            ["Current YOLO dataset", r"external_datasets\super_yolo_dataset\data.yaml"],
            ["Current classification dataset", r"data\merged_dataset_v5\data.yaml"],
            ["Controlled PCA model sweep", r".\.venv311\Scripts\python.exe scripts\pca_feature_model_sweep.py --out runs\ml\pca_feature_model_sweep --components 637 64 128 256 --seed 42"],
            ["Legacy ANN PCA experiment", r".\.venv311\Scripts\python.exe scripts\train_pca_ann.py"],
            ["DL localization rework", r".\.venv311\Scripts\python.exe scripts\classification_to_localization_pipeline.py --max-images 300 --max-visuals 24 --sample-mode stratified --seed 42 --localizer yolo --yolo-conf 0.30 --out-dir runs\dl\localization_rework\yolo_conf030_stratified300_final"],
        ],
        [2200, 7160],
    )

    doc.add_paragraph("8. Artifact Index", style="Heading 1")
    add_table(
        doc,
        ["Area", "Key artifact"],
        [
            ["Workflow summary", r"docs\01_final_report\WORKFLOW_APPROACHES_AND_DL_REWORK.md"],
            ["Current workflow/report notes", r"docs\01_final_report\WORKFLOW_APPROACHES_AND_DL_REWORK.md"],
            ["Newest YOLO dataset", r"external_datasets\super_yolo_dataset"],
            ["Newest classification dataset", r"data\merged_dataset_v5"],
            ["Current ML rerun", r"runs\ml\feature_ml_super_yolo_6class_4k\REPORT.md"],
            ["Legacy ML lecturer run", r"runs\ml\feature_ml_lecturer_6class_4k\REPORT.md"],
            ["PCA model sweep", r"runs\ml\pca_feature_model_sweep\PCA_Model_Sweep_Report.md"],
            ["Legacy ANN PCA report", r"runs\dl\pca_experiments\PCA_Dimensionality_Report.md"],
            ["ML vs DL comparison", r"runs\comparisons\model_comparison\REPORT.md"],
            ["DL architecture comparison", r"runs\dl\comparison_models\model_comparison_report.md"],
            ["DL localization improved", r"runs\dl\localization_rework\yolo_conf030_stratified300_final\REPORT.md"],
        ],
        [2700, 6660],
    )

    doc.add_paragraph("9. Next Actions", style="Heading 1")
    add_numbered(
        doc,
        [
            "Use the newest `super_yolo_dataset` ML rerun for dataset-alignment claims, and label the stronger lecturer run as legacy evidence.",
            "For PCA, cite the controlled Linear SVM row for the about-2% 637-to-128 claim and keep the ANN-only PCA result separate.",
            "Use the localization-first YOLO-localization run at conf=0.30 for final DL localization evidence.",
            "Add final thesis text explaining why classification accuracy and localization metrics are reported separately.",
        ],
    )

    doc.save(OUT_PATH)


if __name__ == "__main__":
    build_doc()
    print(OUT_PATH)
