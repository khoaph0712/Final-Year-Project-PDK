from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent.parent
OUT_DIR = ROOT / "docs" / "01_final_report"
OUT_PATH = OUT_DIR / "WasteWise_Project_Milestones_and_Production_Report.docx"


BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "0B2545"
MUTED = "666666"
LIGHT_GRAY = "F2F4F7"
LIGHT_BLUE_GRAY = "E8EEF5"
CALLOUT = "F4F6F9"
GREEN_FILL = "EAF4EE"
GOLD_FILL = "FFF7DF"
RED_FILL = "FDECEC"
WHITE = "FFFFFF"


LOCAL_EVIDENCE = [
    ("README", r"README.md", "Current high-level project positioning, dataset paths, branch split, and commands."),
    ("Workflow decision", r"docs\01_final_report\WORKFLOW_APPROACHES_AND_DL_REWORK.md", "Explains current ML final branch and DL localization-first classification hierarchical pipeline."),
    ("Project tracker", r"docs\NOTION_PROJECT_TRACKER.md", "Status board, task board, and final report wording."),
    ("Cleanup notes", r"docs\PROJECT_STRUCTURE_AND_CLEANUP.md", "Canonical folders, active datasets, and current result folders."),
    ("Dataset EDA", r"docs\02_dataset_training\DATASET_EDA_AND_TUNING_REPORT.md", "Baseline/tuned dataset cleaning and class-balance evidence."),
    ("External registry", r"docs\02_dataset_training\external_dataset_registry.json", "Downloaded/candidate external datasets, sizes, risks, and actions."),
    ("Source-level dataset analysis", r"docs\02_dataset_training\SOURCE_LEVEL_DATASET_ANALYSIS.md", "Lecturer-directed source-level EDA and per-source ANN/CNN runs."),
    ("Model improvement plan", r"docs\02_dataset_training\MODEL_IMPROVEMENT_PLAN.md", "Hard-case data, classifier retraining, YOLO resume, and Hugging Face deployment notes."),
    ("Current ML rerun", r"runs\ml\feature_ml_super_yolo_6class_4k\REPORT.md", "Newest-dataset 637-feature ML run on super_yolo_dataset."),
    ("Legacy ML run", r"runs\ml\feature_ml_lecturer_6class_4k\REPORT.md", "Lecturer-facing legacy merged_dataset_v3 637-feature ML benchmark."),
    ("PCA sweep", r"runs\ml\pca_feature_model_sweep\PCA_Model_Sweep_Report.md", "Controlled 637-to-PCA component sweep across classical models."),
    ("DL localization sweep", r"runs\dl\localization_rework\THRESHOLD_SWEEP_300.md", "300-image Stage 2 localization threshold comparison."),
    ("Hard-case classifier", r"runs\dl\convnext_hardcase_tuned\RESULT.json", "ConvNeXt hard-case classifier accuracy and macro-F1."),
    ("ConvNeXtV2 material benchmark", r"runs\dl\convnextv2_material_stage1_benchmark\RESULT.json", "ConvNeXtV2 + 637-feature MLP benchmark and latency."),
    ("Stage 0 gate", r"runs\dl\convnextv2_stage0_trash_gate\RESULT.json", "Trash/not-trash/hand/bin gate model metrics."),
    ("YOLO11 vs YOLO26 benchmark", r"runs\detect\yolo11_vs_yolo26_benchmark.json", "Hard-case detector comparison with mAP and latency."),
    ("Web app", r"web\README.md", "Local model API, scanner workflow, and deployment scope."),
]


def rgb(hex_color: str) -> RGBColor:
    return RGBColor.from_string(hex_color)


def set_run_font(run, name: str = "Calibri", size: float | None = None, color: str | None = None, bold: bool | None = None, italic: bool | None = None) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def paragraph_border_bottom(paragraph, color: str = BLUE, size: str = "8", space: str = "4") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = p_bdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        p_bdr.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), space)
    bottom.set(qn("w:color"), color)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, width in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(width))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)
    tbl_header.set(qn("w:val"), "true")


def set_table_width(table, widths_dxa: list[int], indent_dxa: int = 120) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl = table._tbl
    tbl_pr = tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[min(idx, len(widths_dxa) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_cell_text(cell, text: str, bold: bool = False, color: str = "000000", size: float = 9.0, align: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.LEFT) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.10
    run = paragraph.add_run(text)
    set_run_font(run, size=size, color=color, bold=bold)


def style_table(table, widths_dxa: list[int], header_fill: str = LIGHT_GRAY) -> None:
    table.style = "Table Grid"
    table.autofit = False
    set_table_width(table, widths_dxa)
    set_repeat_table_header(table.rows[0])
    for row_i, row in enumerate(table.rows):
        for cell in row.cells:
            if row_i == 0:
                set_cell_shading(cell, header_fill)
            elif row_i % 2 == 0:
                set_cell_shading(cell, "FBFCFD")


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths_dxa: list[int], header_fill: str = LIGHT_GRAY, font_size: float = 8.8) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    for idx, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], header, bold=True, color=INK, size=font_size)
    for row_values in rows:
        row = table.add_row()
        for idx, value in enumerate(row_values):
            text = str(value)
            align = WD_ALIGN_PARAGRAPH.CENTER if text.replace(".", "", 1).replace("%", "", 1).replace(",", "").replace("/", "").replace(" ", "").isdigit() else WD_ALIGN_PARAGRAPH.LEFT
            set_cell_text(row.cells[idx], text, size=font_size, align=align)
    style_table(table, widths_dxa, header_fill=header_fill)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(4)


def add_callout(doc: Document, title: str, body: str, fill: str = CALLOUT, border_color: str = "D7DBE2") -> None:
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, top=130, bottom=130, start=170, end=170)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    set_run_font(r, size=10.5, color=DARK_BLUE, bold=True)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.10
    r2 = p2.add_run(body)
    set_run_font(r2, size=9.6, color="000000")
    style_table(table, [9360], header_fill=fill)
    for row in table.rows:
        for c in row.cells:
            tc_pr = c._tc.get_or_add_tcPr()
            tc_borders = tc_pr.find(qn("w:tcBorders"))
            if tc_borders is None:
                tc_borders = OxmlElement("w:tcBorders")
                tc_pr.append(tc_borders)
            for side in ["top", "left", "bottom", "right"]:
                border = tc_borders.find(qn(f"w:{side}"))
                if border is None:
                    border = OxmlElement(f"w:{side}")
                    tc_borders.append(border)
                border.set(qn("w:val"), "single")
                border.set(qn("w:sz"), "6")
                border.set(qn("w:space"), "0")
                border.set(qn("w:color"), border_color)
    doc.add_paragraph()


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.add_run(item)


def add_figure(doc: Document, path: Path, caption: str, width_in: float = 5.85) -> None:
    if not path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width_in))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(8)
    r = cap.add_run(caption)
    set_run_font(r, size=8.8, color="555555", italic=True)


def add_page_number(paragraph) -> None:
    run = paragraph.add_run("Page ")
    set_run_font(run, size=9, color=MUTED)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def configure_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    title = styles["Title"]
    title.font.name = "Calibri"
    title.font.size = Pt(23)
    title.font.color.rgb = rgb(INK)
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(4)

    subtitle = styles["Subtitle"]
    subtitle.font.name = "Calibri"
    subtitle.font.size = Pt(12.5)
    subtitle.font.color.rgb = rgb("555555")
    subtitle.paragraph_format.space_after = Pt(12)

    for style_name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        st = styles[style_name]
        st.font.name = "Calibri"
        st.font.size = Pt(size)
        st.font.color.rgb = rgb(color)
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True

    for style_name in ["List Bullet", "List Number"]:
        st = styles[style_name]
        st.font.name = "Calibri"
        st.font.size = Pt(11)
        st.paragraph_format.left_indent = Inches(0.5)
        st.paragraph_format.first_line_indent = Inches(-0.25)
        st.paragraph_format.space_after = Pt(8)
        st.paragraph_format.line_spacing = 1.167


def add_header_footer(doc: Document) -> None:
    section = doc.sections[0]
    header_p = section.header.paragraphs[0]
    header_p.text = ""
    header_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = header_p.add_run("WasteWise FYP Milestones Report")
    set_run_font(r, size=9, color=MUTED)

    footer_p = section.footer.paragraphs[0]
    footer_p.text = ""
    footer_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = footer_p.add_run(f"Generated {date.today().isoformat()}  |  ")
    set_run_font(r, size=9, color=MUTED)
    add_page_number(footer_p)


def add_masthead(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("FINAL YEAR PROJECT REPORT")
    set_run_font(r, size=9.5, color=BLUE, bold=True)

    title = doc.add_paragraph(style="Title")
    title.add_run("WasteWise: Project Milestones and Production Build-Up Report")

    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.add_run("From dataset acquisition and feature engineering to ML benchmarks, DL localization, web API, and deployment evidence")

    rows = [
        ("Project", "WasteWise automated waste classification and localization"),
        ("Prepared for", "FYP milestone / final production evidence package"),
        ("Prepared on", date.today().isoformat()),
        ("Workspace", r"C:\FYP"),
        ("Primary evidence", r"docs\, runs\, scripts\, web\, models"),
        ("Report stance", "ML branch is final explainable evidence; DL branch is localization-first classification evidence. Old Grad-CAM classification-first work is retained as experiment history."),
    ]
    for label, value in rows:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r1 = p.add_run(f"{label}: ")
        set_run_font(r1, size=10.5, color="000000", bold=True)
        r2 = p.add_run(value)
        set_run_font(r2, size=10.5, color="000000")

    rule = doc.add_paragraph()
    rule.paragraph_format.space_before = Pt(8)
    paragraph_border_bottom(rule, color=BLUE, size="10", space="3")


def add_exec_summary(doc: Document) -> None:
    add_callout(
        doc,
        "Current project position",
        "WasteWise has matured into a two-branch evidence package. The ML branch is the explainable classification proof using 637 handcrafted features and PCA. The DL branch is the production-facing 2-stage hierarchical pipeline (YOLO localization followed by crop classification verification).",
        fill=GREEN_FILL,
        border_color="B7D8BF",
    )
    add_table(
        doc,
        ["Area", "Final position", "Strongest evidence"],
        [
            ["Dataset foundation", "Current datasets separated by task.", "Classification: `data/merged_dataset_v5` with 29,639 images. Localization: `external_datasets/super_yolo_dataset` with 23,929 images and 102,777 boxes."],
            ["Explainable ML", "Final academic ML branch.", "637-D features; current-dataset XGBoost 0.5408 accuracy / 0.3691 F1; legacy lecturer XGBoost 0.6742 / 0.6506."],
            ["PCA compression", "Valid compact-feature evidence.", "637 to 128 components keeps 99.90% variance; Linear SVM accuracy drops 62.43% to 59.90% (2.52 pp)."],
            ["DL classification", "Classifier/gate and hard-case material model.", "Hard-case ConvNeXt: 93.88% test accuracy / 0.9398 macro F1. ConvNeXtV2+637 MLP: 94.55% test accuracy / 0.9459 macro F1."],
            ["DL localization", "2-stage hierarchical pipeline evidence.", "YOLO conf=0.30 + crop verification on 300 images: precision 0.6999, recall 0.5729, F1 0.6301, mean IoU 0.9057."],
            ["Production", "Web scanner deployed.", "Hugging Face Space update `khoaphung/wastewise-ai` and web API."],
        ],
        [1800, 3100, 4460],
    )


def add_milestones(doc: Document) -> None:
    doc.add_paragraph("1. Chronological Milestones", style="Heading 1")
    doc.add_paragraph(
        "The project moved from raw waste image collection into a final production story through ten practical milestones. Each milestone below names what was built, what was learned, and which local artifact proves it."
    )
    add_table(
        doc,
        ["No.", "Milestone", "Built / learned", "Evidence"],
        [
            ["1", "Problem framing", "Defined WasteWise as waste material recognition plus localization for sorting guidance.", r"README.md; docs\PIPELINE_DIAGRAMS.md"],
            ["2", "Raw data intake", "Collected/registered Roboflow-style merged data, TrashNet-style classification data, TACO-style litter data, RealWaste, Outerview, and hard-case sources.", r"docs\02_dataset_training\external_dataset_registry.json"],
            ["3", "EDA and cleaning", "Found invalid YOLO rows, duplicates, tiny boxes, and split imbalance. Built tuned non-destructive dataset copies.", r"docs\02_dataset_training\DATASET_EDA_AND_TUNING_REPORT.md"],
            ["4", "Source-level audit", "Separated source identity and trained per-source ANN/CNN baselines when standalone classes were valid.", r"docs\02_dataset_training\SOURCE_LEVEL_DATASET_ANALYSIS.md"],
            ["5", "Handcrafted ML", "Built 637-D feature extractor from spatial, FFT, color, and HOG domains, then trained classical models.", r"runs\ml\feature_ml_lecturer_6class_4k; runs\ml\feature_ml_super_yolo_6class_4k"],
            ["6", "PCA study", "Measured how far 637 features could be compressed without losing the report claim.", r"runs\ml\pca_feature_model_sweep\PCA_Model_Sweep_Report.md"],
            ["7", "Legacy DL baselines", "Compared ANN/CNN, MobileNetV2, ResNet50, EfficientNetB0, and old YOLO-first crop verification.", r"runs\dl\comparison_models; runs\comparisons\model_comparison"],
            ["8", "DL rework", "Changed final DL story to 2-stage hierarchical pipeline (localization-first, classification-second).", r"docs\01_final_report\WORKFLOW_APPROACHES_AND_DL_REWORK.md"],
            ["9", "Hard-case improvement", "Downloaded hard-case datasets, retrained material classifier, resumed YOLO, and benchmarked YOLO11 vs YOLO26.", r"docs\02_dataset_training\MODEL_IMPROVEMENT_PLAN.md; runs\detect\yolo11_vs_yolo26_benchmark.json"],
            ["10", "Production surfaces", "Built web scanner/API and deployed Space update.", r"web\README.md; scripts\deploy_hf_space.py"],
        ],
        [600, 1800, 4200, 2760],
    )


def add_datasets(doc: Document) -> None:
    doc.add_paragraph("2. Dataset And Data Number Evidence", style="Heading 1")
    doc.add_paragraph(
        "WasteWise now treats classification and localization datasets separately because each task needs different labels. Classification uses image/crop class support, while localization needs bounding boxes and IoU evaluation."
    )
    add_table(
        doc,
        ["Dataset / source", "Path", "Size", "Role / risk"],
        [
            ["Current classification", r"data\merged_dataset_v5", "29,639 images, 7 classes incl. Background", "Stage 1/classifier evidence; more balanced than older merged split."],
            ["Current localization", r"external_datasets\super_yolo_dataset", "23,929 images, 102,777 boxes, 6 classes", "Stage 2 localization evidence; test support is imbalanced."],
            ["Hard-case classifier", r"data\hard_case_classifier_v1", "16,539 linked records; train 12,600 / val 2,680 / test 2,696 used in classifier result", "Real-world hard-case material classifier."],
            ["YOLO26 hard case", r"external_datasets\yolo26_hardcase_dataset_v1", "20,593 train / 3,450 val / 1,275 test images", "Detector hard-case training/evaluation export."],
            ["TACO official", r"external_datasets\hard_case_full\taco_official", "1,500 images / records downloaded", "Litter detection/segmentation source; sparse class mapping risk."],
            ["RealWaste", r"external_datasets\hard_case_full\realwaste_hf_uci", "4,752 records / images", "Landfill-style classification hard cases."],
            ["Outerview", r"external_datasets\hard_case_full\outerview_global_trash_debris", "25,000 metadata rows, 23,297 extracted images", "OOD review pool; label trust risk."],
        ],
        [1900, 2700, 2700, 2060],
    )
    add_table(
        doc,
        ["Task dataset", "Split / class", "Count"],
        [
            ["merged_dataset_v5", "Train images", "24,039"],
            ["merged_dataset_v5", "Test images", "5,600"],
            ["merged_dataset_v5", "Background total", "4,300"],
            ["merged_dataset_v5", "Cardboard total", "4,225"],
            ["merged_dataset_v5", "Glass total", "4,300"],
            ["merged_dataset_v5", "Metal total", "4,030"],
            ["merged_dataset_v5", "Organic total", "4,300"],
            ["merged_dataset_v5", "Paper total", "4,300"],
            ["merged_dataset_v5", "Plastic total", "4,184"],
            ["super_yolo_dataset", "Images", "23,929"],
            ["super_yolo_dataset", "Boxes", "102,777"],
            ["super_yolo_dataset", "Current ML train crops", "24,000"],
            ["super_yolo_dataset", "Current ML test crops", "2,232"],
        ],
        [2600, 4500, 2260],
    )
    add_callout(
        doc,
        "Dataset caveat that must stay in the final thesis",
        "The newest `super_yolo_dataset` ML test split is not a clean balanced benchmark: glass has only 9 test boxes, cardboard 35, and organic 46. This explains the lower current-dataset macro-F1 and prevents overstating the result.",
        fill=GOLD_FILL,
        border_color="E1C878",
    )


def add_ml(doc: Document) -> None:
    doc.add_paragraph("3. Explainable Machine Learning Branch", style="Heading 1")
    doc.add_paragraph(
        "The ML branch is built for auditability. Every object crop becomes a deterministic handcrafted vector instead of raw pixels. That vector then feeds classical models whose results can be compared with accuracy, F1, confusion matrices, and feature importance."
    )
    add_table(
        doc,
        ["Feature group", "Count", "Reason in this project"],
        [
            ["Spatial / edge", "8", "Brightness, intensity percentiles, gradient mean/std, and edge density."],
            ["Frequency / FFT", "9", "Radial frequency bins plus high-frequency energy for texture and clutter response."],
            ["Color", "44", "HSV histograms and BGR/HSV statistics for material color cues."],
            ["HOG", "576", "Dominant local gradient orientation and shape texture."],
            ["Total", "637", "Fixed explainable vector used by ML, PCA, and hybrid DL work."],
        ],
        [2200, 1100, 6060],
    )
    add_table(
        doc,
        ["Current super_yolo_dataset ML model", "Accuracy", "F1-macro", "Interpretation"],
        [
            ["XGBoost", "0.5408", "0.3691", "Best current-dataset ML result."],
            ["Random Forest", "0.5063", "0.3456", "Tree baseline and feature-importance source."],
            ["ExtraTrees", "0.5045", "0.3414", "Robust noisy-feature baseline."],
            ["Linear SVM", "0.4628", "0.3159", "Scaled linear margin baseline."],
            ["Logistic Regression", "0.4494", "0.3054", "Linear probability baseline."],
            ["Decision Tree", "0.3750", "0.2631", "Simple interpretable threshold baseline."],
        ],
        [3100, 1400, 1400, 3460],
    )
    add_table(
        doc,
        ["Legacy lecturer ML model", "Accuracy", "F1-macro", "Interpretation"],
        [
            ["XGBoost", "0.6742", "0.6506", "Best legacy lecturer-facing result."],
            ["Random Forest", "0.6317", "0.6111", "Useful feature-importance baseline."],
            ["ExtraTrees", "0.6312", "0.6113", "Strong ensemble baseline."],
            ["Linear SVM", "0.5960", "0.5642", "High-dimensional margin baseline."],
            ["Logistic Regression", "0.5864", "0.5558", "Standardized linear baseline."],
            ["Decision Tree", "0.5115", "0.4883", "Low-complexity interpretability baseline."],
        ],
        [2700, 1500, 1500, 3660],
    )
    add_callout(
        doc,
        "Feature-importance result",
        "In the current super_yolo_dataset run, HOG dominates the random-forest feature importance at 59.5808%, followed by color at 29.2090%, frequency at 5.6582%, and spatial features at 5.5520%. This is consistent with waste material differences being driven by shape/texture plus color.",
        fill=GREEN_FILL,
        border_color="B7D8BF",
    )
    add_figure(doc, ROOT / "runs" / "ml" / "feature_ml_super_yolo_6class_4k" / "chart_model_comparison.png", "Figure 1. Current-dataset handcrafted-feature ML model comparison.")
    add_figure(doc, ROOT / "runs" / "ml" / "feature_ml_super_yolo_6class_4k" / "chart_domain_importance.png", "Figure 2. Current-dataset feature-domain importance.")


def add_pca(doc: Document) -> None:
    doc.add_paragraph("4. PCA Compression Evidence", style="Heading 1")
    doc.add_paragraph(
        "PCA was used to test whether the 637 handcrafted features can be compressed for smaller storage and faster inference. The important report rule is to cite the controlled classical-model sweep when claiming the small 637-to-128 loss."
    )
    add_table(
        doc,
        ["Evidence", "Model", "Components", "Explained variance", "Accuracy", "F1-macro", "Drop"],
        [
            ["Controlled ML", "Linear SVM", "637", "100.00%", "62.43%", "0.6235", "0.00 pp"],
            ["Controlled ML", "Linear SVM", "128", "99.90%", "59.90%", "0.5947", "2.52 pp"],
            ["Controlled ML", "LogReg", "637", "100.00%", "60.24%", "0.6019", "0.00 pp"],
            ["Controlled ML", "LogReg", "128", "99.90%", "59.71%", "0.5954", "0.52 pp"],
            ["ANN-only legacy", "MLP", "637", "100.00%", "73.24%", "0.7319 weighted", "0.00 pp"],
            ["ANN-only legacy", "MLP", "128", "99.90%", "68.71%", "0.6863 weighted", "4.53 pp"],
        ],
        [1600, 1400, 1200, 1700, 1200, 1300, 960],
    )
    add_callout(
        doc,
        "Correct thesis wording",
        "Say: 'For Linear SVM in the controlled ML sweep, PCA from 637 to 128 components preserved 99.90% explained variance and reduced accuracy by 2.52 percentage points.' Do not generalize that exact drop to every model.",
        fill=GOLD_FILL,
        border_color="E1C878",
    )
    add_figure(doc, ROOT / "runs" / "ml" / "pca_feature_model_sweep" / "pca_model_sweep_accuracy.png", "Figure 3. PCA component count versus classical-model accuracy.")


def add_dl(doc: Document) -> None:
    doc.add_paragraph("5. Deep Learning Branch And Benchmark Results", style="Heading 1")
    doc.add_paragraph(
        "The DL branch became production-facing after the project separated material classification from localization. Old YOLO-first work stays as experiment history; the final DL story uses a classification/gate stage and then a localization-only stage."
    )
    add_table(
        doc,
        ["Experiment", "Metric result", "Decision"],
        [
            ["ML vs tiny CNN baseline", "Tiny CNN accuracy 0.4813 / F1 0.4303; ML ExtraTrees and RF around 0.631 accuracy.", "ML stayed the primary explainable branch."],
            ["Architecture comparison", "EfficientNetB0 94.29%; ResNet50 89.76%; MobileNetV2 85.43%.", "EfficientNetB0 justified as compact high-accuracy classifier evidence."],
            ["Hard-case ConvNeXt", "Test accuracy 93.88%; macro F1 0.9398.", "Promote after local API smoke tests."],
            ["ConvNeXtV2 + 637 MLP", "Test accuracy 94.55%; macro F1 0.9459; 1,405-D combined feature vector.", "Best material-stage benchmark currently present."],
            ["Stage 0 trash gate", "Test accuracy 97.81%; macro F1 0.9777 across trash/not_trash/hand/bin.", "Useful front-door filter before material routing."],
        ],
        [2600, 3600, 3160],
    )
    add_table(
        doc,
        ["Deep model", "Accuracy", "Size", "CPU latency", "Use"],
        [
            ["MobileNetV2", "85.43%", "20.07 MB", "253.5 ms", "Lightweight baseline."],
            ["ResNet50", "89.76%", "161.52 MB", "163.1 ms", "Large baseline; high size cost."],
            ["EfficientNetB0", "94.29%", "29.21 MB", "288.6 ms", "Primary historical classifier choice."],
            ["ConvNeXt hard-case", "93.88%", "PyTorch .pth", "not listed here", "Current hard-case tuned classifier."],
            ["ConvNeXtV2 + 637 MLP", "94.55%", "joblib head", "68.53 ms mean CPU", "Current strongest material benchmark."],
        ],
        [2300, 1200, 1600, 1800, 2460],
    )
    add_table(
        doc,
        ["YOLO detector benchmark", "mAP50", "mAP50-95", "Precision", "Recall", "Mean latency"],
        [
            ["old_yolo11", "0.6693", "0.4961", "0.7447", "0.5930", "50.14 ms"],
            ["new_yolo26n", "0.6723", "0.5054", "0.7569", "0.5922", "52.11 ms"],
        ],
        [2500, 1300, 1400, 1400, 1200, 1560],
    )
    add_figure(doc, ROOT / "runs" / "dl" / "comparison_models" / "confusion_matrix_grid.png", "Figure 4. DL architecture confusion-matrix grid.")


def add_localization(doc: Document) -> None:
    doc.add_paragraph("6. Final DL Workflow: Classification First, Localization Second", style="Heading 1")
    doc.add_paragraph(
        "The current final DL workflow answers two questions separately: first, what class evidence is visible; second, where object evidence appears. YOLO is used as the Stage 2 box localizer, not as the final material classifier."
    )
    add_table(
        doc,
        ["Stage", "Implementation", "Metric"],
        [
            ["Stage 0", "Trash-state gate: trash, not_trash, hand, bin.", "Accuracy / macro F1."],
            ["Stage 1", "Material classifier/gate using ConvNeXt/EfficientNet and handcrafted features.", "Accuracy, macro F1, calibration, class confusion."],
            ["Stage 2", "YOLO localization-only module or Grad-CAM baseline.", "Precision, recall, F1, matched IoU, TP/FP/FN."],
            ["Decision logic", "Review gate for uncertain or out-of-distribution scans.", "Wrong confident route reduction."],
        ],
        [1500, 4860, 3000],
    )
    add_table(
        doc,
        ["Stage 2 localizer", "Images", "GT boxes", "Pred boxes", "TP", "FP", "FN", "Precision", "Recall", "F1", "Mean IoU"],
        [
            ["Grad-CAM baseline", "60", "261", "74", "19", "55", "242", "0.2568", "0.0728", "0.1134", "0.7127"],
            ["YOLO conf=0.25", "60", "261", "233", "148", "85", "113", "0.6352", "0.5670", "0.5991", "0.9012"],
            ["YOLO conf=0.30", "300", "1,152", "943", "660", "283", "492", "0.6999", "0.5729", "0.6301", "0.9057"],
            ["YOLO conf=0.35", "300", "1,152", "815", "617", "198", "535", "0.7571", "0.5356", "0.6274", "0.9043"],
            ["YOLO conf=0.40", "300", "1,152", "738", "593", "145", "559", "0.8035", "0.5148", "0.6275", "0.9050"],
        ],
        [1900, 700, 800, 900, 550, 550, 550, 900, 800, 700, 1010],
        font_size=7.6,
    )
    add_callout(
        doc,
        "Promoted localization setting",
        "`--localizer yolo --yolo-conf 0.30` is promoted because it gives the best F1 and recall among the 300-image threshold checks while keeping mean matched IoU above 0.90.",
        fill=GREEN_FILL,
        border_color="B7D8BF",
    )
    add_figure(
        doc,
        ROOT / "runs" / "dl" / "localization_rework" / "yolo_conf030_stratified300_final" / "visuals" / "rf_garbage_metal391_jpg.rf.d2d79150c42df8cd64bea8d65acc58ab_yolo.jpg",
        "Figure 5. Example localization-first crop-verification output.",
    )


def add_production(doc: Document) -> None:
    doc.add_paragraph("7. Production Build-Up", style="Heading 1")
    doc.add_paragraph(
        "Production is represented by a web/API scanner for public demonstration. The web scanner is the primary production path because it is tied to local PyTorch/YOLO model artifacts and Hugging Face Space deployment notes."
    )
    add_table(
        doc,
        ["Surface", "Built", "Current production note"],
        [
            ["Web scanner", "Plain HTML/CSS/JS frontend with Python model API in `web/server.py`.", "Uses uploads, no static demo samples; saves browser History; local model API falls back only when unavailable."],
            ["Backend models", "ConvNeXt + 637-feature classifier, scaler, YOLO detector, crop verification, review gate.", "YOLO settings restored to `imgsz=960`, `max_det=80`, crop verification 80."],
            ["Hugging Face Space", "Space `khoaphung/wastewise-ai` updated.", "Review-gate update deployed at commit `7a8af57ec3c3efc5fa468d08c0f79c5e5b1f8f46`."],
        ],
        [2100, 3900, 3360],
    )
    add_numbered(
        doc,
        [
            "Local production smoke path: run `web/server.py --port 4178`, upload real images, and verify material, confidence, route, waste state, and review flag.",
            "Deployment path: run `scripts/deploy_hf_space.py` after replacing stable model artifacts and checking `/api/health`.",
        ],
    )


def add_sources(doc: Document) -> None:
    doc.add_paragraph("8. Literature, Books, Datasets, And Technical References", style="Heading 1")
    doc.add_paragraph(
        "This section links the local implementation choices to primary papers, dataset pages, official docs, and reference books. These are not substitutes for local run metrics; they explain why the selected methods are technically relevant."
    )
    add_table(
        doc,
        ["Reference", "Why it matters to WasteWise", "URL"],
        [
            ["YOLO original paper", "Single-stage real-time object detection foundation.", "https://arxiv.org/abs/1506.02640"],
            ["Ultralytics YOLO11 docs", "Localizer family used in the project and cited as software/docs rather than a formal YOLO11 paper.", "https://docs.ultralytics.com/models/yolo11/"],
            ["Ultralytics YOLO26 docs", "New detector family benchmarked locally against YOLO11.", "https://docs.ultralytics.com/models/yolo26/"],
            ["ConvNeXt", "Modern ConvNet backbone family used for material classifier direction.", "https://arxiv.org/abs/2201.03545"],
            ["EfficientNet", "Architecture-comparison baseline and historical tuned classifier.", "https://arxiv.org/abs/1905.11946"],
            ["ResNet", "Deep residual baseline in architecture comparison.", "https://arxiv.org/abs/1512.03385"],
            ["MobileNet", "Mobile/edge classification baseline.", "https://arxiv.org/abs/1704.04861"],
            ["HOG", "576 of 637 handcrafted ML features come from HOG texture/shape descriptors.", "https://lear.inrialpes.fr/people/triggs/pubs/Dalal-cvpr05.pdf"],
            ["SIFT", "Classical local-feature context for handcrafted vision descriptors.", "https://www.cs.ubc.ca/~lowe/papers/ijcv04.pdf"],
            ["TACO", "Real-world litter detection/segmentation source and YOLO hard-case source.", "https://arxiv.org/abs/2003.06975"],
            ["TrashNet", "Classic clean-background waste classification reference dataset.", "https://github.com/garythung/trashnet"],
            ["RealWaste", "Real landfill-style hard-case classification source.", "https://archive.ics.uci.edu/dataset/908/realwaste"],
            ["ZeroWaste", "Industrial cluttered waste detection/segmentation context.", "https://arxiv.org/abs/2106.02740"],
            ["Hugging Face Docker Spaces", "Deployment surface for web/API production demo.", "https://huggingface.co/docs/hub/spaces-sdks-docker"],
            ["Ultralytics TFLite export", "Model export path relevant to mobile/edge deployment.", "https://docs.ultralytics.com/integrations/tflite/"],
            ["Szeliski CV book", "General computer vision theory reference.", "https://szeliski.org/Book/"],
            ["Deep Learning book", "General deep learning theory reference.", "https://www.deeplearningbook.org/"],
            ["TinyML book", "Reference for mobile/edge ML constraints.", "https://www.oreilly.com/library/view/tinyml/9781492052036/"],
        ],
        [1900, 3900, 3560],
        font_size=7.6,
    )
    doc.add_paragraph("9. Local Artifact Evidence Index", style="Heading 1")
    add_table(doc, ["Artifact", "Path", "Use"], [list(row) for row in LOCAL_EVIDENCE], [1800, 3300, 4260], font_size=7.6)


def add_final_actions(doc: Document) -> None:
    doc.add_paragraph("10. Final Production Checklist", style="Heading 1")
    add_table(
        doc,
        ["Checklist item", "Status", "Evidence / next proof"],
        [
            ["Dataset registry complete", "Mostly done", "Current registry names full hard-case downloads and candidate datasets."],
            ["ML explainability branch", "Done", "Feature extraction, model comparison, PCA, confusion matrices, feature importance."],
            ["DL classifier benchmark", "Done", "Hard-case ConvNeXt and ConvNeXtV2+637 MLP benchmark artifacts."],
            ["DL localization benchmark", "Done", "300-image threshold sweep with promoted conf=0.30."],
            ["YOLO26 promotion decision", "Measured, not automatically promoted", "YOLO26n improves mAP slightly but mean CPU latency is slightly slower than old YOLO11 in the local 80-image benchmark."],
            ["Web deployment", "Done / needs final smoke before demo", "Hugging Face Space update recorded; run local API and 10 real uploads before presentation."],
            ["Large artifact storage", "Open", "Keep model binaries outside Git; use release, Drive, or Hugging Face artifact workflow."],
        ],
        [2700, 1900, 4760],
    )
    add_callout(
        doc,
        "Final report wording",
        "Use this concise position: 'WasteWise final evidence separates explainable ML classification from DL localization. The ML branch reports accuracy, macro-F1, PCA, and feature importance. The DL branch reports localization precision, recall, F1, and matched IoU for the 2-stage hierarchical pipeline. The alternative classification-first Grad-CAM pipeline is retained as experiment comparison.'",
        fill=GREEN_FILL,
        border_color="B7D8BF",
    )


def build_doc() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_styles(doc)
    add_header_footer(doc)
    add_masthead(doc)
    add_exec_summary(doc)
    add_milestones(doc)
    add_datasets(doc)
    add_ml(doc)
    add_pca(doc)
    add_dl(doc)
    add_localization(doc)
    add_production(doc)
    add_sources(doc)
    add_final_actions(doc)
    doc.save(OUT_PATH)


if __name__ == "__main__":
    build_doc()
    print(OUT_PATH)
